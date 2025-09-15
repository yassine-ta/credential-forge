"""Word format synthesizer using agent-generated content."""

import random
from pathlib import Path
from typing import Dict, Any

from .format_synthesizer import FormatSynthesizer
from ..utils.exceptions import SynthesizerError

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordFormatSynthesizer(FormatSynthesizer):
    """Word format synthesizer that structures agent-generated content."""
    
    def __init__(self, output_dir: str, format_type: str = 'docx', ultra_fast_mode: bool = False):
        """Initialize Word format synthesizer.
        
        Args:
            output_dir: Output directory for generated files
            format_type: Word format type (docx, doc, docm, rtf)
            ultra_fast_mode: Enable ultra-fast mode with minimal validation
        """
        super().__init__(output_dir, ultra_fast_mode)
        self.format_type = format_type
    
    def synthesize(self, content_structure: Dict[str, Any]) -> str:
        """Structure content into Word format.
        
        Args:
            content_structure: Generated content from agents
            
        Returns:
            Path to generated Word file
        """
        try:
            # Validate content structure
            self._validate_content_structure(content_structure)
            
            # Embed credentials in content
            content_structure = self._embed_credentials_in_content(content_structure)
            
            # Generate filename and save
            filename = self._generate_filename(content_structure)
            file_path = self._get_file_path(filename)
            
            if DOCX_AVAILABLE and self.format_type in ['docx', 'docm']:
                # Create Word document with python-docx
                self._create_word_with_docx(content_structure, file_path)
            else:
                # Create simple text-based document
                self._create_simple_document(content_structure, file_path)
            
            # Log stats
            self._log_generation_stats(content_structure)
            
            return str(file_path)
            
        except Exception as e:
            self.generation_stats['errors'] += 1
            raise SynthesizerError(f"Word synthesis failed: {e}")
    
    def _create_word_with_docx(self, content_structure: Dict[str, Any], file_path: Path) -> None:
        """Create Word document using python-docx."""
        doc = Document()
        
        # Title
        title = doc.add_heading(content_structure.get('title', 'Document'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        metadata = content_structure.get('metadata', {})
        if metadata:
            doc.add_paragraph(f"Topic: {metadata.get('topic', 'N/A')}")
            doc.add_paragraph(f"Language: {content_structure.get('language', 'en')}")
            doc.add_paragraph(f"Format: {content_structure.get('format_type', 'unknown')}")
            doc.add_paragraph("")  # Empty line
        
        # Sections
        sections = content_structure.get('sections', [])
        for section in sections:
            section_title = section.get('title', 'Section')
            section_content = section.get('content', '')
            
            # Section heading
            heading = doc.add_heading(section_title, level=1)
            
            # Section content
            content_paragraph = doc.add_paragraph()
            content_paragraph.add_run(section_content)
            
            # Add some spacing
            doc.add_paragraph("")
        
        # Save document
        doc.save(str(file_path))
    
    def _create_simple_document(self, content_structure: Dict[str, Any], file_path: Path) -> None:
        """Create simple text-based document."""
        content = f"""
{content_structure.get('title', 'Document')}
{'=' * len(content_structure.get('title', 'Document'))}

"""
        
        # Metadata
        metadata = content_structure.get('metadata', {})
        if metadata:
            content += f"Topic: {metadata.get('topic', 'N/A')}\n"
            content += f"Language: {content_structure.get('language', 'en')}\n"
            content += f"Format: {content_structure.get('format_type', 'unknown')}\n\n"
        
        # Sections
        sections = content_structure.get('sections', [])
        for section in sections:
            section_title = section.get('title', 'Section')
            section_content = section.get('content', '')
            
            content += f"""
{section_title}
{'=' * len(section_title)}

{section_content}

"""
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    def _generate_filename(self, content_structure: Dict[str, Any]) -> str:
        """Generate Word filename."""
        title = content_structure.get('title', 'document')
        timestamp = self._get_current_timestamp()
        random_id = random.randint(1000, 9999)
        
        # Clean title for filename
        clean_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        clean_title = clean_title.replace(' ', '_').lower()
        
        return f"document_{clean_title}_{timestamp}_{random_id}.{self.format_type}"
    
    def _get_current_timestamp(self) -> str:
        """Get current timestamp for filename."""
        from datetime import datetime
        return datetime.now().strftime('%Y%m%d_%H%M%S')


class RTFFormatSynthesizer(WordFormatSynthesizer):
    """RTF format synthesizer that structures agent-generated content."""
    
    def __init__(self, output_dir: str, ultra_fast_mode: bool = False):
        """Initialize RTF format synthesizer.
        
        Args:
            output_dir: Output directory for generated files
            ultra_fast_mode: Enable ultra-fast mode with minimal validation
        """
        super().__init__(output_dir, 'rtf', ultra_fast_mode)
    
    def synthesize(self, content_structure: Dict[str, Any]) -> str:
        """Structure content into RTF format.
        
        Args:
            content_structure: Generated content from agents
            
        Returns:
            Path to generated RTF file
        """
        try:
            # Validate content structure
            self._validate_content_structure(content_structure)
            
            # Embed credentials in content
            content_structure = self._embed_credentials_in_content(content_structure)
            
            # Generate filename and save
            filename = self._generate_filename(content_structure)
            file_path = self._get_file_path(filename)
            
            # Create RTF document
            self._create_rtf_document(content_structure, file_path)
            
            # Log stats
            self._log_generation_stats(content_structure)
            
            return str(file_path)
            
        except Exception as e:
            self.generation_stats['errors'] += 1
            raise SynthesizerError(f"RTF synthesis failed: {e}")
    
    def _create_rtf_document(self, content_structure: Dict[str, Any], file_path: Path) -> None:
        """Create RTF document."""
        rtf_content = r"{\rtf1\ansi\deff0"
        
        # Title
        title = content_structure.get('title', 'Document')
        rtf_content += f"\\fs24\\b {title}\\b0\\fs20\\par\\par"
        
        # Metadata
        metadata = content_structure.get('metadata', {})
        if metadata:
            rtf_content += f"Topic: {metadata.get('topic', 'N/A')}\\par"
            rtf_content += f"Language: {content_structure.get('language', 'en')}\\par"
            rtf_content += f"Format: {content_structure.get('format_type', 'unknown')}\\par\\par"
        
        # Sections
        sections = content_structure.get('sections', [])
        for section in sections:
            section_title = section.get('title', 'Section')
            section_content = section.get('content', '')
            
            rtf_content += f"\\fs18\\b {section_title}\\b0\\fs20\\par"
            rtf_content += f"{section_content}\\par\\par"
        
        rtf_content += "}"
        
        # Write to file
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
